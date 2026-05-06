from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/mac/Documents/New project/agri-platform/docs")
OUTPUT_DOCX = ROOT / "ESP32智能灌溉系统交付文档-20260504.docx"
FUNCTION_MD = ROOT / "ESP32智能灌溉系统功能说明.md"
MANUAL_MD = ROOT / "ESP32智能灌溉系统操作手册.md"


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_bottom_border(paragraph, color="D9E2F2", size="6", space="1") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def set_table_layout(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def set_run_font(run, name="Arial", size=12, bold=False, color: str | None = None) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    title = styles["Title"]
    title.font.name = "Arial"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string("17324D")

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Arial"
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = RGBColor.from_string("5B6B7F")

    for style_name, size in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)):
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10 if style_name != "Heading 1" else 12)
        style.paragraph_format.space_after = Pt(4 if style_name != "Heading 1" else 6)

    if "Code Block" not in styles:
        code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        code_style.base_style = styles["Normal"]
        code_style.font.name = "Courier New"
        code_style._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
        code_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
        code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        code_style.font.size = Pt(10.5)
        code_style.paragraph_format.space_before = Pt(3)
        code_style.paragraph_format.space_after = Pt(3)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.text = ""
    run = p.add_run("ESP32 智能灌溉系统交付文档")
    set_run_font(run, size=9, color="617286")
    set_paragraph_bottom_border(p)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.text = ""
    run = p.add_run("第 ")
    set_run_font(run, size=9, color="617286")
    add_page_number(p)
    run = p.add_run(" 页")
    set_run_font(run, size=9, color="617286")


def clean_markdown_text(text: str) -> str:
    def replace_link(match: re.Match[str]) -> str:
        label = match.group(1)
        target = match.group(2)
        if label == target:
            return label
        return f"{label}（{target}）"

    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", replace_link, text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = text.replace("**", "")
    return text


def add_cover(document: Document) -> None:
    p = document.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("ESP32 智能灌溉系统交付文档")

    p = document.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("功能说明与操作手册")

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("适用对象：soil-001 / A-PUMP-001")
    set_run_font(run, size=11, color="4E5D6C")

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("版本日期：2026-05-04")
    set_run_font(run, size=11, color="4E5D6C")

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("当前固件：1.0.4-safety")
    set_run_font(run, size=11, color="4E5D6C")

    document.add_paragraph("")

    table = document.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    set_table_layout(table)

    rows = [
        ("项目名称", "Agri Platform 智能灌溉现场系统"),
        ("主控设备", "ESP32 DevKit 32E"),
        ("现场网关", "soil-001"),
        ("执行器", "A-PUMP-001"),
        ("平台地址", "http://82.156.45.208/")
    ]
    for row, (k, v) in zip(table.rows, rows):
        row.cells[0].width = Inches(1.7)
        row.cells[1].width = Inches(4.8)
        row.cells[0].text = k
        row.cells[1].text = v
        for idx, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                for run in paragraph.runs:
                    set_run_font(run, size=10.5, bold=(idx == 0), color="213547")
            set_cell_border(
                cell,
                top={"val": "single", "sz": 6, "color": "D9E2F2"},
                bottom={"val": "single", "sz": 6, "color": "D9E2F2"},
                left={"val": "single", "sz": 6, "color": "D9E2F2"},
                right={"val": "single", "sz": 6, "color": "D9E2F2"},
            )
        set_cell_shading(row.cells[0], "EEF4FB")

    document.add_paragraph("")
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(
        "本文档汇总当前现场系统功能、操作步骤、OTA 升级链路、本地自治策略和默认参数，供交付、培训与后续维护使用。"
    )
    set_run_font(run, size=11, color="4E5D6C")

    document.add_page_break()


def add_markdown_content(document: Document, title: str, markdown_text: str) -> None:
    heading = document.add_paragraph(title, style="Heading 1")
    heading.paragraph_format.space_before = Pt(0)

    lines = markdown_text.splitlines()
    in_code_block = False
    code_lines: list[str] = []

    for raw_line in lines:
        line = raw_line.rstrip()

        if line.startswith("```"):
            if in_code_block:
                for code_line in code_lines:
                    p = document.add_paragraph(style="Code Block")
                    p.paragraph_format.left_indent = Inches(0.2)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.space_before = Pt(0)
                    run = p.add_run(code_line)
                    set_run_font(run, name="Courier New", size=10.5, color="22313F")
                code_lines = []
                in_code_block = False
                document.add_paragraph("")
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        if not line.strip():
            continue

        if line.startswith("# "):
            continue
        if line.startswith("## "):
            document.add_paragraph(clean_markdown_text(line[3:]), style="Heading 2")
            continue
        if line.startswith("### "):
            document.add_paragraph(clean_markdown_text(line[4:]), style="Heading 3")
            continue

        if re.match(r"^\d+\.\s+", line):
            prefix = re.match(r"^(\d+\.)\s+", line).group(1)
            text = clean_markdown_text(re.sub(r"^\d+\.\s+", "", line))
            p = document.add_paragraph(style="Normal")
            run = p.add_run(f"{prefix} {text}")
            set_run_font(run, size=12, color="22313F")
            continue

        if line.startswith("- "):
            text = clean_markdown_text(line[2:])
            p = document.add_paragraph(style="List Bullet")
            p.add_run(text)
            continue

        p = document.add_paragraph(style="Normal")
        p.add_run(clean_markdown_text(line))


def add_section_break(document: Document) -> None:
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.top_margin = Inches(1)
    new_section.bottom_margin = Inches(1)
    new_section.left_margin = Inches(1)
    new_section.right_margin = Inches(1)
    new_section.header.is_linked_to_previous = True
    new_section.footer.is_linked_to_previous = True


def main() -> None:
    document = Document()
    configure_document(document)
    add_cover(document)

    function_md = FUNCTION_MD.read_text(encoding="utf-8")
    manual_md = MANUAL_MD.read_text(encoding="utf-8")

    add_markdown_content(document, "第一部分 功能说明", function_md)
    add_section_break(document)
    add_markdown_content(document, "第二部分 操作手册", manual_md)

    document.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
